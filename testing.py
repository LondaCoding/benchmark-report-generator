from docx import Document

# Create a new Document
doc = Document("Template - Copy.docx")

# Add a Title
doc.add_heading('Document Title', level=1)
doc.add_paragraph('This is a paragraph.')
doc.add_heading('Document Title', level=2)
doc.add_paragraph('This is another paragraph.')

title3 = doc.add_heading('Document Title', level=3)
title3.paragraph_format.space_after = None
doc.add_paragraph('This is another paragraph.')

# Add Paragraphs

# Add a Table
table = doc.add_table(rows=3, cols=3)
for i in range(3):
    for j in range(3):
        cell = table.cell(i, j)
        cell.text = f'Row {i+1}, Col {j+1}'

# Save the Document
doc.save('Template - Copy.docx')

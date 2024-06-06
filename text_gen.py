from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from openai import OpenAI
import time
import os
import re

client= OpenAI()
# gpt_counter= 0
error_log= []
asset_counter= 0
found_fields_counter= 0
added_field_counter= 0
comment_counter= 0
accepted_types= {'PDP':'Desarrolladas produciendo', 
               'PNP':'Desarrolladas no produciendo', 
               'PND':'No desarrolladas',
               'PRB':'Probables',
               'PS':'Posibles'}

def createComments():
    arr = [['Reservas Netas'],
           ['Regalias'],
           ['Revenue'],
           ['OPEX Fijo'],
           ['OPEX Semifijo'],
           ['OPEX Variable'],
           ['OPEX Ecopetrol'],
           ['CAPEX Total'],
           ['Flujo de Caja Operativo'],
           ['Abandono']]
    return arr


def retreiveDocumentInfo(file_path):
    worksheets= ['BM', 'BM (2)', 'BM (3)', 'BM (4)']
    found= False
    for worksheet in worksheets:
        try:
            wb = load_workbook(file_path)
            ws = wb[worksheet]
            print(f'Opened worksheet {worksheet}')
            found= True
            break
        except Exception as e:
            # error= f'Worksheet "{worksheet}" is not present.'
            # print(error)
            None

    if not found:
        error= f'None of the worksheets are pressent: {worksheets}'
        return(error)
    
    # Iterate over all cells with comments
    current_reserve_type = ws['B1'].value
    comments = createComments()
    
    document_text = []
    for col in ws.iter_cols():
        
        #verify if category has changed
        if col[0].value:
            if not(col[0].value==current_reserve_type) and not (col[0].value=="Variable"):
                #print(col[0].value)
                there_are_comments= False
                for variable in comments:
                    if len(variable) > 1:
                        there_are_comments= True
                if there_are_comments:
                    document_text.append(current_reserve_type)
                    document_text.append(comments)
                    comments= createComments()
                current_reserve_type= col[0].value
            
        #traverse the column
        for cell in col[:33]:
            if cell.comment:
                
                expression = re.compile(r'Comment:\n(.*)', re.DOTALL)
                comment = expression.search(cell.comment.text)
                if not comment:
                    expression = re.compile(r'Comentario:\n(.*)', re.DOTALL)
                    comment = expression.search(cell.comment.text)
                comment= comment.group(1)
                comments[(cell.row-3)//3].append(comment)

    #documment has finished, add the last type of reserve  
    document_text.append(current_reserve_type)
    document_text.append(comments)
    return document_text        


def findBenchmark(field_path):
    #single file documentation
    directory_items= os.listdir(field_path)
    benchmarks= [item for item in directory_items if "_New_" not in item]
    if len(benchmarks) < 1:
        error_message= 'Error: there was no benchmark found (1).'
        return error_message
    else: 
        None

    #find latest modified benchmark
    most_recent_date= datetime(1900, 1, 1)
    most_recent_file : str = None
    for file in benchmarks:
        new_file_date= datetime.fromtimestamp(os.path.getmtime(os.path.join(field_path, file)))
        if new_file_date > most_recent_date:
            most_recent_date= new_file_date
            most_recent_file= file

    if not most_recent_file:    
        error_message= 'Error: there was no benchmark found (2)'
        return error_message
    else:
        excel_path= os.path.join(field_path, most_recent_file)
    
        # Check the extension and replace if necessary
        # base_name, extension = os.path.splitext(excel_path)
        # if extension == '.xlsx':
        #    print('Error: the extension of the file is .xlsx ')
            
        # elif extension == '.xlsm':
        #     None
        # else:
        #     raise ValueError("File must have either a .xlsx or .xlsm extension")
        # print('modified file extention:', excel_path)

        return excel_path


def addFieldToDocument(doc:Document, field_info, field_name):    
    global comment_counter
    print_type= True
    for location_type in field_info:
        #skip if not an accepted type type
        if print_type==False:
            print_type=True
            continue
        #it's a header 1
        if type(location_type) is str:
            if location_type in accepted_types:
                title_text= accepted_types[location_type]+f' ({location_type})'
                doc.add_heading(title_text, level=3)
            else:
                print_type= False
                print(f'The type: "{location_type}" was skipped')
        #it's a variable array
        else:
            #Get variables that don't havecomments
            no_comment_variables= []
            for variable in location_type:
                if len(variable) <= 1:
                    no_comment_variables.append(variable[0])
                    location_type.remove(variable)
            no_comment_variables_title : str = ''
            for variable in no_comment_variables:
                no_comment_variables_title+= variable + ', '
            no_comment_variables_title= no_comment_variables_title[:-2]
            doc.add_heading(no_comment_variables_title, level=4)
            doc.add_paragraph('Calculo OK')

            #add comments
            for variable in location_type:
                if len(variable) > 1:
                    doc.add_heading(variable[0], level=4)
                    paragraph= ''
                    for comment in variable[1:]:
                        comment_counter+= 1
                        paragraph+= comment.strip(' \t\n\r')
                    #CHATGPT CALL
                    # paragraph= aiCorrection(field_name, title_text, variable[0], paragraph)
                    para= doc.add_paragraph(paragraph)
                    para.alignment= 3
            doc.add_paragraph('')            
    doc.add_paragraph('')    
            
    return doc

# def aiCorrection(field_name, location_type, variable, text):
#     model= "gpt-3.5-turbo"
#     system_context= '''Eres un proveedor de software para empresas de hidrocarburos. 
#     Has comparado las variables calculadas por tu software Planning Space (abreviado PS) 
#     con los valores brindados por un auditor y realizaste textos que explican las diferencias 
#     entre los valores del auditor y "Planning Space". Se te brindara: "Nombre del campo", 
#     "Tipo de campo", "Variable comparada", "Texto a corregir". Tu mision es devolver unicamente el texto de 
#     mensaje (sin comillas o texto adicional), con ortografía y sintaxis corregida a la 
#     perfección y redactado en tercera persona, teniendo en cuenta el contexto del mensaje.'''
#     content= f'Nombre del campo: {field_name}\nTipo de campo: {location_type}\nVariable comparada: {variable}\nTexto a corregir: {text}'

#     print(f'Calling "{model}" for location type "{location_type}", variable "{variable}"')
#     completion= client.chat.completions.create(
#         model=model,
#         messages=[
#             {"role":"system", "content":system_context},
#             {"role":"user", "content":content}
#         ]
#     )
#     global gpt_counter
#     gpt_counter+= 1
#     return completion.choices[0].message.content

def traverseAsset(doc, asset_path, asset_name):
    #Get the folders in the directory
    asset_items= os.listdir(asset_path)
    fields= [item for item in asset_items if os.path.isdir(os.path.join(asset_path, item))]
        
    #add asset name to doc
    print('Name of the asset:', asset_name)
    global asset_counter
    last_paragraph= doc.add_heading(f'{asset_counter+1}. {asset_name}', level=1)
    last_paragraph.alignment = 1

    global added_field_counter
    #There's only one field in the asset
    if len(fields) < 1:
        print("Theres only one field in the asset")
        temporal_doc= createField(doc, asset_path, asset_name, asset_name)
        if temporal_doc:
            doc=temporal_doc 
            added_field_counter+= 1
        else:
            print('The field info was not added to the document')
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.clear()
            return doc
    #there are multiple fields in the asset
    else:
        print(f'Fields of the asset "{asset_name}":', fields)
        field_counter= 0
        for field in fields:
            doc.add_heading(f'{asset_counter+1}.{field_counter+1}. {field}', level=2)
            field_path= os.path.join(asset_path, field)
            temporal_doc= createField(doc, field_path, field, asset_name)
            if temporal_doc:
                doc=temporal_doc 
                field_counter+= 1
                added_field_counter+= 1
            else:
                print('The field info was not added to the document')
                last_paragraph = doc.paragraphs[-1]
                last_paragraph._element.getparent().remove(last_paragraph._element)
                continue

    asset_counter+= 1
    print(f'Asset "{asset_name}" added to doccument.')
    return doc
    
  
def createField(doc, field_path, field, asset):
    print()
    print('Creating field:', field)
    print('Finding benchmark...')
    excel_path= findBenchmark(field_path)
    
    if 'Error:' in excel_path:
        error_message= f'Couldnt find the benchmark for the only field of "{field}" of asset: "{asset}": {excel_path}'
        error_log.append(error_message)
        print(error_message)
        return None
    global found_fields_counter
    found_fields_counter+= 1

    print(f'Fetching field "{field}" comments...')
    excel_content= retreiveDocumentInfo(excel_path)
    if type(excel_content) is str:
        error_message= f'Couldnt retrieve field "{field}" from asset "{asset}". Error: {excel_content}'
        error_log.append(error_message)
        print(error_message)
        return None
    
    print(f'Adding field "{field}" comments...')
    return addFieldToDocument(doc, excel_content, field)


def generateReportFolder():
    #create directory where reports will be saved
    contributors= [item for item in os.listdir(os.getcwd()) if os.path.isdir(item)]
    try:
        contributors.remove("Reportes Generados")
        print('"Reportes Generados" not considered as an asset')
    except ValueError:
        None
    try:
        contributors.remove(".git")
        print('".git" not considered as an asset')
    except ValueError:
        None
    try:
        contributors.remove("Test")
        print('"Test" not considered as an asset')
    except ValueError:
        None

    print('Contributors:', contributors)
    print()
    
    template_path= os.path.join(os.getcwd(), "Template.docx") #change to Template!!!!!!!!!!
    doc= Document(template_path)
    doc.paragraphs[16].alignment= 1

    for contributer in contributors:
        contributer_dir= os.path.join(os.getcwd(), contributer)
        assets= [item for item in os.listdir(contributer_dir) if os.path.isdir(os.path.join(contributer_dir, item))]
        
        for asset in assets:
            #modify the title & it's style
            regex= r'\d+. (.*)'
            try:
                asset_name= re.findall(regex, asset)[0]
            except IndexError as e:
                print(f'The folder numeration <#. > finished, there are no more finished assets. Program finished.')
                print('Program finished')
                continue

            print(asset_name)
            asset_path= os.path.join(contributer_dir, asset)

            print('CREATING ASSET:', asset)
            doc= traverseAsset(doc, asset_path, asset_name)
            print()
            print()
    
    #modify title
    doc.paragraphs[16].text= 'Reporte de Valores por Activo'
    target_paragraph = doc.paragraphs[16] 
    for run in target_paragraph.runs:
        run.font.size = Pt(26)  # Change to the desired font size in points
    for run in target_paragraph.runs:
        run.font.color.rgb = RGBColor(0, 112, 192)  # Change RGB values for the desired color
    for run in target_paragraph.runs:
        run.font.name = 'Calibri'  # Change to the desired font name

    #Save file
    new_file_location= os.path.join(os.getcwd(), f'Reporte.docx')
    doc.save(new_file_location)
    print(f'Document saved in "{new_file_location}"')


    print('ASSET COUNTER:', asset_counter)
    print('COMMENT COUNTER:', comment_counter)
    print('FOLDER FIELDS FOUND:', found_fields_counter)
    print('ADDED FIELDS:', added_field_counter)
    # print('CHATGPT REQUEST COUNTER:', gpt_counter)
    print('\nERRORS DURING EXECUTION:')
    for error in error_log:
        print(error)
        print()


generateReportFolder()